"""
Forensic CPA AI - Document Analysis & Processing System

A Flask-based web application for forensic accounting document analysis.
Supports PDF, Excel, and Word document processing with AI-powered insights.
"""

from flask import Flask, render_template, request, jsonify, send_file
import pdfplumber
import pandas as pd
from openpyxl import load_workbook
from docx import Document
import os
from werkzeug.utils import secure_filename
from datetime import datetime
import json

app = Flask(__name__)

# Configuration
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'xlsx', 'xls', 'docx', 'doc'}

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


def allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def process_pdf(file_path):
    """Extract text and tables from PDF."""
    results = {
        'text': [],
        'tables': [],
        'metadata': {}
    }

    try:
        with pdfplumber.open(file_path) as pdf:
            results['metadata'] = {
                'pages': len(pdf.pages),
                'creator': pdf.metadata.get('Creator', 'Unknown'),
                'producer': pdf.metadata.get('Producer', 'Unknown')
            }

            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text
                text = page.extract_text()
                if text:
                    results['text'].append({
                        'page': page_num,
                        'content': text
                    })

                # Extract tables
                tables = page.extract_tables()
                if tables:
                    for table_num, table in enumerate(tables, 1):
                        results['tables'].append({
                            'page': page_num,
                            'table': table_num,
                            'data': table
                        })

    except Exception as e:
        results['error'] = str(e)

    return results


def process_excel(file_path):
    """Extract data from Excel file."""
    results = {
        'sheets': [],
        'metadata': {}
    }

    try:
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        results['metadata']['sheet_names'] = excel_file.sheet_names
        results['metadata']['sheet_count'] = len(excel_file.sheet_names)

        # Load workbook for additional metadata
        wb = load_workbook(file_path, read_only=True)

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            sheet_data = {
                'name': sheet_name,
                'rows': len(df),
                'columns': list(df.columns),
                'preview': df.head(10).to_dict('records'),
                'summary': {
                    'numeric_columns': df.select_dtypes(include=['number']).columns.tolist(),
                    'text_columns': df.select_dtypes(include=['object']).columns.tolist(),
                    'date_columns': df.select_dtypes(include=['datetime']).columns.tolist()
                }
            }

            # Add statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                sheet_data['statistics'] = df[numeric_cols].describe().to_dict()

            results['sheets'].append(sheet_data)

    except Exception as e:
        results['error'] = str(e)

    return results


def process_word(file_path):
    """Extract text and tables from Word document."""
    results = {
        'paragraphs': [],
        'tables': [],
        'metadata': {}
    }

    try:
        doc = Document(file_path)

        results['metadata'] = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables)
        }

        # Extract paragraphs
        for para_num, para in enumerate(doc.paragraphs, 1):
            if para.text.strip():
                results['paragraphs'].append({
                    'number': para_num,
                    'text': para.text,
                    'style': para.style.name
                })

        # Extract tables
        for table_num, table in enumerate(doc.tables, 1):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)

            results['tables'].append({
                'number': table_num,
                'rows': len(table.rows),
                'columns': len(table.columns),
                'data': table_data
            })

    except Exception as e:
        results['error'] = str(e)

    return results


@app.route('/')
def index():
    """Home page."""
    return jsonify({
        'application': 'Forensic CPA AI',
        'version': '1.0.0',
        'description': 'Document Analysis & Processing System for Forensic Accounting',
        'endpoints': {
            '/': 'API information',
            '/ui': 'Web interface for document upload',
            '/health': 'Health check',
            '/api/upload': 'Upload and process documents (POST)',
            '/api/analyze/<file_id>': 'Get analysis results',
            '/api/files': 'List uploaded files'
        },
        'supported_formats': ['PDF', 'Excel (.xlsx, .xls)', 'Word (.docx, .doc)']
    })


@app.route('/ui')
def ui():
    """Web interface for document upload."""
    return render_template('index.html')


@app.route('/health')
def health():
    """Health check endpoint."""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'service': 'Forensic CPA AI'
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Upload and process a document."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not allowed'}), 400

    # Save file
    filename = secure_filename(file.filename)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    file_id = f"{timestamp}_{filename}"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_id)
    file.save(file_path)

    # Process based on file type
    ext = filename.rsplit('.', 1)[1].lower()

    if ext == 'pdf':
        results = process_pdf(file_path)
    elif ext in ['xlsx', 'xls']:
        results = process_excel(file_path)
    elif ext in ['docx', 'doc']:
        results = process_word(file_path)
    else:
        return jsonify({'error': 'Unsupported file type'}), 400

    # Save results
    results_path = file_path + '.json'
    with open(results_path, 'w') as f:
        json.dump(results, f, indent=2)

    return jsonify({
        'success': True,
        'file_id': file_id,
        'filename': filename,
        'file_type': ext,
        'processed_at': datetime.now().isoformat(),
        'results': results
    })


@app.route('/api/analyze/<file_id>')
def analyze_file(file_id):
    """Get analysis results for a specific file."""
    results_path = os.path.join(app.config['UPLOAD_FOLDER'], file_id + '.json')

    if not os.path.exists(results_path):
        return jsonify({'error': 'File not found or not processed'}), 404

    with open(results_path, 'r') as f:
        results = json.load(f)

    return jsonify(results)


@app.route('/api/files')
def list_files():
    """List all uploaded and processed files."""
    files = []

    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
        if not filename.endswith('.json'):
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            stat = os.stat(file_path)

            files.append({
                'file_id': filename,
                'filename': filename.split('_', 1)[1] if '_' in filename else filename,
                'size': stat.st_size,
                'uploaded_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'has_results': os.path.exists(file_path + '.json')
            })

    return jsonify({
        'count': len(files),
        'files': sorted(files, key=lambda x: x['uploaded_at'], reverse=True)
    })


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_DEBUG', 'False').lower() == 'true'

    print(f"""
    ╔══════════════════════════════════════════════════════════════╗
    ║          Forensic CPA AI - Document Processing System        ║
    ╚══════════════════════════════════════════════════════════════╝

    🚀 Server starting on http://localhost:{port}

    📄 Supported file types: PDF, Excel, Word
    🔍 Ready to process forensic accounting documents

    API Endpoints:
      GET  /              - API information
      GET  /health        - Health check
      POST /api/upload    - Upload & process documents
      GET  /api/analyze/<file_id> - Get analysis results
      GET  /api/files     - List uploaded files

    Press CTRL+C to stop the server
    """)

    app.run(host='0.0.0.0', port=port, debug=debug)
